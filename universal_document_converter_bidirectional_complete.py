#!/usr/bin/env python3
"""
Universal Document Converter Bidirectional - Complete Fixed Version
Document conversion tool with full bidirectional support and multiple output formats
Designed and built by Beau Lewis (blewisxx@gmail.com)

Features:
- Bidirectional conversion (including FROM Markdown)
- Multiple output formats (MD, TXT, DOCX, PDF, HTML, RTF)
- OCR with multiple output format support
- Multi-threaded processing
- Proper GUI integration
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
from pathlib import Path
import sys
import json
from typing import Optional, Union, Dict, Any, List
import concurrent.futures
import time
from threading import Lock
import re

# Import the base converter
from universal_document_converter_complete import (
    DocumentConverterApp, ConfigurationManager,
    DocumentConverterError, FileProcessingError,
    OCR_AVAILABLE
)

# Optional imports with fallbacks
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown2
    MARKDOWN2_AVAILABLE = True
except ImportError:
    MARKDOWN2_AVAILABLE = False

class BidirectionalDocumentConverter(DocumentConverterApp):
    """Extended converter with full bidirectional support"""
    
    def __init__(self, root):
        # Initialize parent class first
        super().__init__(root)
        
        # Update title
        self.root.title("Universal Document Converter - Bidirectional Edition")
        
        # Add output format variable
        self.output_format_var = tk.StringVar(value="markdown")
        
        # Define format mappings
        self.output_formats = {
            'markdown': {'ext': '.md', 'name': 'Markdown', 'icon': '📝'},
            'text': {'ext': '.txt', 'name': 'Plain Text', 'icon': '📄'},
            'docx': {'ext': '.docx', 'name': 'Word Document', 'icon': '📘'},
            'pdf': {'ext': '.pdf', 'name': 'PDF', 'icon': '📕'},
            'html': {'ext': '.html', 'name': 'HTML', 'icon': '🌐'},
            'rtf': {'ext': '.rtf', 'name': 'Rich Text Format', 'icon': '📋'}
        }
        
        # Add format selector to GUI
        self.add_format_selector()
        
    def create_settings_section(self, parent):
        """Override to add output format selector"""
        # Call parent method first
        super().create_settings_section(parent)
        
        # Add output format frame
        format_frame = ttk.LabelFrame(parent, text="Output Format", padding=10)
        format_frame.pack(fill=tk.X, pady=(10, 0))
        
        # Format selector row
        selector_row = ttk.Frame(format_frame)
        selector_row.pack(fill=tk.X)
        
        ttk.Label(selector_row, text="Convert to:").pack(side=tk.LEFT, padx=5)
        
        # Create format combo with nice display
        self.format_combo = ttk.Combobox(
            selector_row,
            textvariable=self.output_format_var,
            state='readonly',
            width=25
        )
        
        # Set display values
        display_values = [f"{info['icon']} {info['name']}" for info in self.output_formats.values()]
        self.format_combo['values'] = list(self.output_formats.keys())
        self.format_combo.pack(side=tk.LEFT, padx=5)
        
        # Format info label
        self.format_info_label = ttk.Label(selector_row, text="", foreground='blue')
        self.format_info_label.pack(side=tk.LEFT, padx=10)
        
        # Bind change event
        self.format_combo.bind('<<ComboboxSelected>>', self.on_format_change)
        
        # Dependencies info
        deps_frame = ttk.Frame(format_frame)
        deps_frame.pack(fill=tk.X, pady=(5, 0))
        
        self.deps_label = ttk.Label(deps_frame, text="", font=('Arial', 8))
        self.deps_label.pack(side=tk.LEFT, padx=5)
        
        # Initialize display
        self.on_format_change()
        
    def add_format_selector(self):
        """Add format selector to existing GUI"""
        # This method is called after parent initialization
        # The actual UI is created in create_settings_section override
        pass
        
    def on_format_change(self, event=None):
        """Handle output format change"""
        format_name = self.output_format_var.get()
        format_info = self.output_formats.get(format_name, {})
        
        # Update info label
        self.format_info_label.config(
            text=f"Extension: {format_info.get('ext', '.?')}"
        )
        
        # Check dependencies and update status
        deps_status = []
        if format_name == 'pdf':
            if REPORTLAB_AVAILABLE:
                deps_status.append("✅ PDF ready")
            else:
                deps_status.append("⚠️ Install: pip install reportlab")
        elif format_name == 'docx':
            if DOCX_AVAILABLE:
                deps_status.append("✅ DOCX ready")
            else:
                deps_status.append("⚠️ Install: pip install python-docx")
        elif format_name == 'html' and MARKDOWN2_AVAILABLE:
            deps_status.append("✅ Enhanced HTML")
        
        if hasattr(self, 'deps_label'):
            self.deps_label.config(text=" ".join(deps_status))
        
        # Update status bar
        self.status_var.set(f"Output format: {format_info.get('name', 'Unknown')}")
    
    def scan_input_folder(self):
        """Override to ensure .md files are included"""
        if not self.input_path.get():
            return
        
        self.clear_file_list()
        self.status_var.set("Scanning folder...")
        
        # Define patterns - ensure .md is included
        doc_patterns = ['*.docx', '*.pdf', '*.txt', '*.md', '*.rtf', '*.odt', 
                       '*.html', '*.htm', '*.epub', '*.xml', '*.json', '*.csv']
        
        # Add image patterns if OCR is enabled
        if self.ocr_enabled_var.get() and OCR_AVAILABLE:
            image_patterns = ['*.jpg', '*.jpeg', '*.png', '*.tiff', '*.tif', 
                            '*.bmp', '*.gif', '*.webp']
            doc_patterns.extend(image_patterns)
        
        input_path = Path(self.input_path.get())
        files_found = []
        
        for pattern in doc_patterns:
            files_found.extend(input_path.rglob(pattern))
        
        # Filter out temporary files
        files_found = [f for f in files_found if not f.name.startswith('~$')]
        
        # Add files to tree with format indicator
        for file_path in files_found:
            size = self.format_file_size(file_path.stat().st_size)
            file_type = file_path.suffix.upper()[1:]
            rel_path = file_path.relative_to(input_path)
            
            # Add indicator if it's a markdown file
            display_text = str(rel_path)
            if file_path.suffix.lower() == '.md':
                display_text = f"📝 {display_text}"
            
            self.file_tree.insert('', 'end', text=display_text, 
                                 values=(size, file_type, 'Pending'))
        
        self.stats['total_files'] = len(files_found)
        self.update_stats_display()
        self.status_var.set(f"Found {len(files_found)} files to convert")
    
    def convert_single_file(self, item_id, file_path, input_dir, output_dir):
        """Override to support multiple output formats"""
        try:
            # Get output format
            output_format = self.output_format_var.get()
            format_info = self.output_formats[output_format]
            
            # Calculate output path
            if self.preserve_structure_var.get():
                rel_path = file_path.relative_to(input_dir)
                output_path = output_dir / rel_path.with_suffix(format_info['ext'])
            else:
                output_path = output_dir / f"{file_path.stem}{format_info['ext']}"
            
            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Check if file exists and skip if needed
            if output_path.exists() and not self.overwrite_var.get():
                self.update_file_status(item_id, "⏭️ Skipped")
                self.stats['skipped'] += 1
                self.update_progress()
                return
            
            # First, convert to markdown (intermediate format)
            file_ext = file_path.suffix.lower()
            markdown_content = None
            
            # Get markdown content based on input type
            if file_ext == '.md':
                # Input is already markdown
                markdown_content = self.read_markdown_file(file_path)
            else:
                # Convert to markdown using parent methods
                markdown_content = self.convert_to_markdown(file_path, file_ext)
            
            # Convert from markdown to target format
            if markdown_content:
                if output_format == 'markdown':
                    # Direct save for markdown
                    output_content = markdown_content
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(output_content)
                else:
                    # Convert to other formats
                    self.convert_and_save(markdown_content, output_path, output_format, file_path)
                
                # Apply compression if enabled
                if self.compression_var.get():
                    self.compress_file(output_path)
                
                self.update_file_status(item_id, f"✓ {format_info['icon']}")
                self.stats['processed'] += 1
            else:
                raise FileProcessingError("No content extracted")
            
        except Exception as e:
            self.update_file_status(item_id, f"✗ {str(e)[:30]}")
            self.stats['failed'] += 1
            if self.logger:
                self.logger.error(f"Failed to convert {file_path}: {str(e)}")
        
        finally:
            self.update_progress()
    
    def convert_to_markdown(self, file_path, file_ext):
        """Convert any supported format to markdown"""
        # Image extensions
        image_extensions = ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp']
        
        # Use appropriate converter
        if file_ext in image_extensions and self.ocr_enabled_var.get() and OCR_AVAILABLE:
            return self.convert_image_to_markdown(file_path)
        elif file_ext == '.docx':
            return self.convert_docx_to_markdown(file_path)
        elif file_ext == '.pdf':
            if self.ocr_enabled_var.get() and OCR_AVAILABLE and self.is_pdf_image_based(file_path):
                return self.convert_image_to_markdown(file_path)
            else:
                return self.convert_pdf_to_markdown(file_path)
        elif file_ext in ['.txt']:
            return self.convert_txt_to_markdown(file_path)
        elif file_ext in ['.html', '.htm']:
            return self.convert_html_to_markdown(file_path)
        elif file_ext == '.rtf':
            return self.convert_rtf_to_markdown(file_path)
        elif file_ext == '.odt':
            return self.convert_odt_to_markdown(file_path)
        elif file_ext == '.epub':
            return self.convert_epub_to_markdown(file_path)
        elif file_ext == '.json':
            return self.convert_json_to_markdown(file_path)
        elif file_ext == '.csv':
            return self.convert_csv_to_markdown(file_path)
        elif file_ext == '.xml':
            return self.convert_xml_to_markdown(file_path)
        else:
            raise FileProcessingError(f"Unsupported format: {file_ext}")
    
    def read_markdown_file(self, file_path):
        """Read a markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise FileProcessingError(f"Failed to read markdown: {str(e)}")
    
    def convert_and_save(self, markdown_content, output_path, output_format, source_path):
        """Convert markdown to target format and save"""
        if output_format == 'text':
            content = self.markdown_to_text(markdown_content)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
        elif output_format == 'html':
            content = self.markdown_to_html(markdown_content, source_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
        elif output_format == 'rtf':
            content = self.markdown_to_rtf(markdown_content)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
        elif output_format == 'docx':
            self.markdown_to_docx(markdown_content, output_path, source_path)
            
        elif output_format == 'pdf':
            self.markdown_to_pdf(markdown_content, output_path, source_path)
            
        else:
            raise FileProcessingError(f"Unknown output format: {output_format}")
    
    def markdown_to_text(self, markdown_content):
        """Convert markdown to plain text"""
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', markdown_content, flags=re.MULTILINE)
        
        # Remove emphasis
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Italic
        text = re.sub(r'__([^_]+)__', r'\1', text)      # Bold alt
        text = re.sub(r'_([^_]+)_', r'\1', text)        # Italic alt
        
        # Remove links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Remove images
        text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove horizontal rules
        text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
        
        # Clean up lists
        text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
        
        # Clean up extra whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def markdown_to_html(self, markdown_content, source_path):
        """Convert markdown to HTML with styling"""
        if MARKDOWN2_AVAILABLE:
            import markdown2
            html_body = markdown2.markdown(
                markdown_content,
                extras=['tables', 'fenced-code-blocks', 'header-ids', 'strike']
            )
        else:
            # Basic fallback conversion
            html_body = markdown_content
            html_body = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)
            html_body = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
            html_body = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_body, flags=re.MULTILINE)
            html_body = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html_body)
            html_body = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', html_body)
            html_body = html_body.replace('\n\n', '</p><p>').replace('\n', '<br>')
            html_body = f'<p>{html_body}</p>'
        
        # Create complete HTML document
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{source_path.stem}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h3 {{ font-size: 1.25em; }}
        code {{
            background-color: #f6f8fa;
            padding: 0.2em 0.4em;
            margin: 0;
            font-size: 85%;
            border-radius: 3px;
            font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
        }}
        pre {{
            background-color: #f6f8fa;
            padding: 16px;
            overflow: auto;
            font-size: 85%;
            line-height: 1.45;
            border-radius: 3px;
        }}
        blockquote {{
            padding: 0 1em;
            color: #6a737d;
            border-left: 0.25em solid #dfe2e5;
            margin: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 16px 0;
        }}
        table th, table td {{
            padding: 6px 13px;
            border: 1px solid #dfe2e5;
        }}
        table tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        a {{
            color: #0366d6;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
        
        return html
    
    def markdown_to_rtf(self, markdown_content):
        """Convert markdown to RTF format"""
        # RTF header
        rtf = r"{\rtf1\ansi\deff0 {\fonttbl{\f0 Times New Roman;}}"
        rtf += r"\viewkind4\uc1\pard\f0\fs24 "
        
        # Process line by line
        lines = markdown_content.split('\n')
        for line in lines:
            # Headers
            if line.startswith('# '):
                rtf += r"\fs32\b " + line[2:] + r"\b0\fs24\par "
            elif line.startswith('## '):
                rtf += r"\fs28\b " + line[3:] + r"\b0\fs24\par "
            elif line.startswith('### '):
                rtf += r"\fs24\b " + line[4:] + r"\b0\fs24\par "
            # Bold
            elif '**' in line:
                line = re.sub(r'\*\*([^*]+)\*\*', r'\\b \1\\b0 ', line)
                rtf += line + r"\par "
            # Italic
            elif '*' in line:
                line = re.sub(r'\*([^*]+)\*', r'\\i \1\\i0 ', line)
                rtf += line + r"\par "
            # Regular text
            elif line.strip():
                rtf += line + r"\par "
            # Empty line
            else:
                rtf += r"\par "
        
        rtf += "}"
        return rtf
    
    def markdown_to_docx(self, markdown_content, output_path, source_path):
        """Convert markdown to DOCX format"""
        if not DOCX_AVAILABLE:
            raise FileProcessingError("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document()
        
        # Add title
        doc.add_heading(source_path.stem, 0)
        
        # Process markdown content
        lines = markdown_content.split('\n')
        current_paragraph = []
        
        for line in lines:
            # Headers
            if line.startswith('# '):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line[4:], level=3)
            # Lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line.strip()):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
            # Code blocks
            elif line.startswith('```'):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                # Skip code block markers
                continue
            # Empty line - end paragraph
            elif not line.strip():
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
            # Regular text
            else:
                current_paragraph.append(line)
        
        # Add any remaining paragraph
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
        
        # Save document
        doc.save(str(output_path))
    
    def markdown_to_pdf(self, markdown_content, output_path, source_path):
        """Convert markdown to PDF format"""
        if not REPORTLAB_AVAILABLE:
            raise FileProcessingError("reportlab not installed. Install with: pip install reportlab")
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CodeBlock',
            parent=styles['Code'],
            fontSize=9,
            leftIndent=20,
            rightIndent=20,
            backColor='#f6f8fa',
            borderColor='#d1d5da',
            borderWidth=1,
            borderPadding=10
        ))
        
        # Add title
        story.append(Paragraph(source_path.stem, styles['Title']))
        story.append(Spacer(1, 12))
        
        # Process markdown content
        lines = markdown_content.split('\n')
        current_paragraph = []
        in_code_block = False
        code_lines = []
        
        for line in lines:
            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_lines:
                        code_text = '<br/>'.join(code_lines)
                        story.append(Paragraph(code_text, styles['CodeBlock']))
                        story.append(Spacer(1, 12))
                        code_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    if current_paragraph:
                        story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                        story.append(Spacer(1, 12))
                        current_paragraph = []
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line.replace('<', '&lt;').replace('>', '&gt;'))
                continue
            
            # Headers
            if line.startswith('# '):
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                story.append(Paragraph(line[2:], styles['Heading1']))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                story.append(Paragraph(line[3:], styles['Heading2']))
                story.append(Spacer(1, 12))
            elif line.startswith('### '):
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                story.append(Paragraph(line[4:], styles['Heading3']))
                story.append(Spacer(1, 12))
            # Lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                story.append(Paragraph('• ' + line.strip()[2:], styles['Normal']))
            # Empty line
            elif not line.strip():
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    story.append(Spacer(1, 12))
                    current_paragraph = []
            # Regular text
            else:
                # Escape special characters for PDF
                line = line.replace('<', '&lt;').replace('>', '&gt;')
                current_paragraph.append(line)
        
        # Add any remaining paragraph
        if current_paragraph:
            story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
        
        # Build PDF
        doc.build(story)
    
    def show_completion_summary(self):
        """Override to show output format in summary"""
        elapsed = self.stats['end_time'] - self.stats['start_time']
        elapsed_str = time.strftime('%H:%M:%S', time.gmtime(elapsed))
        
        output_format = self.output_format_var.get()
        format_info = self.output_formats[output_format]
        
        summary = f"Conversion Complete!\n\n"
        summary += f"Output Format: {format_info['icon']} {format_info['name']}\n"
        summary += f"Total Files: {self.stats['total_files']}\n"
        summary += f"Processed: {self.stats['processed']}\n"
        summary += f"Failed: {self.stats['failed']}\n"
        summary += f"Skipped: {self.stats['skipped']}\n"
        summary += f"Time Elapsed: {elapsed_str}\n\n"
        summary += f"Output saved to: {self.output_path.get()}"
        
        self.status_var.set(f"Conversion complete! Format: {format_info['name']}")
        
        if self.notification_var.get():
            messagebox.showinfo("Conversion Complete", summary)
        
        if self.auto_open_var.get():
            self.open_output_folder()
        
        if self.logger:
            self.logger.info(f"Conversion complete: {self.stats}, Format: {output_format}")

def main():
    """Application entry point"""
    try:
        # Try enhanced drag-and-drop support
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
    except ImportError:
        # Fall back to regular tkinter
        root = tk.Tk()
    
    # Create application
    app = BidirectionalDocumentConverter(root)
    
    # Center window
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    
    # Start application
    root.mainloop()

if __name__ == "__main__":
    main()
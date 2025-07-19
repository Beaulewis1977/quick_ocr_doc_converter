#!/usr/bin/env python3
"""
Universal Document Converter Bidirectional - Enterprise Edition
Complete document conversion tool with bidirectional support and multiple output formats
Supports converting FROM and TO multiple formats including Markdown
Designed and built by Beau Lewis (blewisxx@gmail.com)

Features:
- Bidirectional conversion (including FROM Markdown)
- Multiple output formats (MD, TXT, DOCX, PDF, HTML, RTF)
- OCR with multiple output format support
- Multi-threaded processing
- Format auto-detection
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
from pathlib import Path
import sys
import json
from typing import Optional, Union, Dict, Any, List
import concurrent.futures
import time
from threading import Lock

# Import the base converter
from universal_document_converter_complete import (
    DocumentConverterApp, ConfigurationManager, 
    DocumentConverterError, FileProcessingError,
    OCR_AVAILABLE
)

# Additional imports for output formats
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown2
    MARKDOWN2_AVAILABLE = True
except ImportError:
    MARKDOWN2_AVAILABLE = False

class BidirectionalDocumentConverter(DocumentConverterApp):
    """Extended converter with bidirectional support and multiple output formats"""
    
    def __init__(self, root):
        # Initialize parent class
        super().__init__(root)
        
        # Update title
        self.root.title("Universal Document Converter - Bidirectional Edition")
        
        # Add output format variable
        self.output_format_var = tk.StringVar(value="markdown")
        
        # Update supported formats to include markdown as input
        self.input_formats = {
            'documents': ['*.docx', '*.pdf', '*.txt', '*.md', '*.rtf', '*.odt', 
                         '*.html', '*.htm', '*.epub', '*.xml', '*.json', '*.csv'],
            'images': ['*.jpg', '*.jpeg', '*.png', '*.tiff', '*.tif', '*.bmp', 
                      '*.gif', '*.webp']
        }
        
        self.output_formats = {
            'markdown': {'ext': '.md', 'name': 'Markdown'},
            'text': {'ext': '.txt', 'name': 'Plain Text'},
            'docx': {'ext': '.docx', 'name': 'Word Document'},
            'pdf': {'ext': '.pdf', 'name': 'PDF'},
            'html': {'ext': '.html', 'name': 'HTML'},
            'rtf': {'ext': '.rtf', 'name': 'Rich Text Format'}
        }
        
        # Add output format selector to GUI
        self.add_output_format_selector()
    
    def add_output_format_selector(self):
        """Add output format selector to the settings section"""
        # Find the settings frame and add format selector
        # This would be added to the create_settings_section method
        pass
    
    def setup_ui(self):
        """Override to add output format selector"""
        super().setup_ui()
        
        # Add output format selector after calling parent setup_ui
        # Find the general settings tab and add the selector
        for child in self.root.winfo_children():
            if isinstance(child, ttk.Frame):
                self.add_format_selector_to_frame(child)
                break
    
    def add_format_selector_to_frame(self, parent_frame):
        """Add format selector to the appropriate frame"""
        # This is a simplified version - in production would integrate properly
        format_frame = ttk.LabelFrame(parent_frame, text="Output Format", padding=10)
        format_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(format_frame, text="Convert to:").pack(side=tk.LEFT, padx=5)
        
        format_combo = ttk.Combobox(
            format_frame,
            textvariable=self.output_format_var,
            values=list(self.output_formats.keys()),
            state='readonly',
            width=20
        )
        format_combo.pack(side=tk.LEFT, padx=5)
        
        # Bind change event
        format_combo.bind('<<ComboboxSelected>>', self.on_format_change)
    
    def on_format_change(self, event=None):
        """Handle output format change"""
        format_name = self.output_format_var.get()
        format_info = self.output_formats.get(format_name, {})
        
        # Update status
        self.status_var.set(f"Output format: {format_info.get('name', 'Unknown')}")
        
        # Check if required libraries are available
        if format_name == 'pdf' and not REPORTLAB_AVAILABLE:
            messagebox.showwarning(
                "PDF Support", 
                "PDF output requires 'reportlab' package.\nInstall with: pip install reportlab"
            )
        elif format_name == 'docx' and not DOCX_AVAILABLE:
            messagebox.showwarning(
                "DOCX Support",
                "DOCX output requires 'python-docx' package.\nInstall with: pip install python-docx"
            )
    
    def scan_input_folder(self):
        """Override to include .md files"""
        if not self.input_path.get():
            return
        
        self.clear_file_list()
        self.status_var.set("Scanning folder...")
        
        # Get all patterns including .md
        all_patterns = self.input_formats['documents'].copy()
        
        # Add image patterns if OCR is enabled
        if self.ocr_enabled_var.get() and OCR_AVAILABLE:
            all_patterns.extend(self.input_formats['images'])
        
        input_path = Path(self.input_path.get())
        files_found = []
        
        for pattern in all_patterns:
            files_found.extend(input_path.rglob(pattern))
        
        # Filter out temporary files
        files_found = [f for f in files_found if not f.name.startswith('~$')]
        
        # Add files to tree
        for file_path in files_found:
            size = self.format_file_size(file_path.stat().st_size)
            file_type = file_path.suffix.upper()[1:]
            rel_path = file_path.relative_to(input_path)
            
            self.file_tree.insert('', 'end', text=str(rel_path), 
                                 values=(size, file_type, 'Pending'))
        
        self.stats['total_files'] = len(files_found)
        self.update_stats_display()
        self.status_var.set(f"Found {len(files_found)} files to convert")
    
    def convert_single_file(self, item_id, file_path, input_dir, output_dir):
        """Override to support multiple output formats"""
        try:
            # Get output format
            output_format = self.output_format_var.get()
            format_info = self.output_formats[output_format]
            
            # Calculate output path
            if self.preserve_structure_var.get():
                rel_path = file_path.relative_to(input_dir)
                output_path = output_dir / rel_path.with_suffix(format_info['ext'])
            else:
                output_path = output_dir / f"{file_path.stem}{format_info['ext']}"
            
            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Check if file exists and skip if needed
            if output_path.exists() and not self.overwrite_var.get():
                self.update_file_status(item_id, "Skipped")
                self.stats['skipped'] += 1
                self.update_progress()
                return
            
            # First, convert to markdown (intermediate format)
            file_ext = file_path.suffix.lower()
            markdown_content = None
            
            # Check if it's an image file for OCR
            image_extensions = ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp']
            
            if file_ext in image_extensions and self.ocr_enabled_var.get() and OCR_AVAILABLE:
                markdown_content = self.convert_image_to_markdown(file_path)
            elif file_ext == '.docx':
                markdown_content = self.convert_docx_to_markdown(file_path)
            elif file_ext == '.pdf':
                if self.ocr_enabled_var.get() and OCR_AVAILABLE and self.is_pdf_image_based(file_path):
                    markdown_content = self.convert_image_to_markdown(file_path)
                else:
                    markdown_content = self.convert_pdf_to_markdown(file_path)
            elif file_ext in ['.txt']:
                markdown_content = self.convert_txt_to_markdown(file_path)
            elif file_ext in ['.md']:
                # If input is markdown, just read it
                markdown_content = self.read_markdown_file(file_path)
            elif file_ext in ['.html', '.htm']:
                markdown_content = self.convert_html_to_markdown(file_path)
            elif file_ext == '.rtf':
                markdown_content = self.convert_rtf_to_markdown(file_path)
            elif file_ext == '.json':
                markdown_content = self.convert_json_to_markdown(file_path)
            elif file_ext == '.csv':
                markdown_content = self.convert_csv_to_markdown(file_path)
            elif file_ext == '.xml':
                markdown_content = self.convert_xml_to_markdown(file_path)
            else:
                raise FileProcessingError(f"Unsupported input format: {file_ext}")
            
            # Now convert from markdown to target format
            if markdown_content:
                output_content = self.convert_from_markdown(markdown_content, output_format, file_path)
                
                # Save the output
                self.save_output_file(output_path, output_content, output_format)
                
                # Apply compression if enabled
                if self.compression_var.get():
                    self.compress_file(output_path)
                
                self.update_file_status(item_id, "✓ Complete")
                self.stats['processed'] += 1
            else:
                raise FileProcessingError("No content extracted")
            
        except Exception as e:
            self.update_file_status(item_id, f"✗ Failed: {str(e)}")
            self.stats['failed'] += 1
            if self.logger:
                self.logger.error(f"Failed to convert {file_path}: {str(e)}")
        
        finally:
            self.update_progress()
    
    def read_markdown_file(self, file_path):
        """Read a markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise FileProcessingError(f"Failed to read markdown file: {str(e)}")
    
    def convert_from_markdown(self, markdown_content, output_format, source_path):
        """Convert markdown content to specified output format"""
        if output_format == 'markdown':
            return markdown_content
        elif output_format == 'text':
            return self.markdown_to_text(markdown_content)
        elif output_format == 'html':
            return self.markdown_to_html(markdown_content)
        elif output_format == 'docx':
            return self.markdown_to_docx(markdown_content, source_path)
        elif output_format == 'pdf':
            return self.markdown_to_pdf(markdown_content, source_path)
        elif output_format == 'rtf':
            return self.markdown_to_rtf(markdown_content)
        else:
            raise FileProcessingError(f"Unsupported output format: {output_format}")
    
    def markdown_to_text(self, markdown_content):
        """Convert markdown to plain text"""
        # Remove markdown formatting
        import re
        
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', markdown_content, flags=re.MULTILINE)
        
        # Remove bold and italic
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # Remove links
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove images
        text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
        
        # Remove horizontal rules
        text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
        
        # Clean up extra whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def markdown_to_html(self, markdown_content):
        """Convert markdown to HTML"""
        if MARKDOWN2_AVAILABLE:
            import markdown2
            html = markdown2.markdown(
                markdown_content,
                extras=['tables', 'fenced-code-blocks', 'header-ids']
            )
            
            # Wrap in basic HTML structure
            return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
        else:
            # Fallback - basic conversion
            html_content = markdown_content.replace('\n', '<br>\n')
            return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Document</title>
</head>
<body>
{html_content}
</body>
</html>"""
    
    def markdown_to_docx(self, markdown_content, source_path):
        """Convert markdown to DOCX - returns None, saves directly"""
        if not DOCX_AVAILABLE:
            raise FileProcessingError("python-docx not installed")
        
        # This is a placeholder - actual implementation would parse markdown
        # and create a properly formatted DOCX
        return None  # Indicates direct file writing in save_output_file
    
    def markdown_to_pdf(self, markdown_content, source_path):
        """Convert markdown to PDF - returns None, saves directly"""
        if not REPORTLAB_AVAILABLE:
            raise FileProcessingError("reportlab not installed")
        
        # This is a placeholder - actual implementation would parse markdown
        # and create a properly formatted PDF
        return None  # Indicates direct file writing in save_output_file
    
    def markdown_to_rtf(self, markdown_content):
        """Convert markdown to RTF"""
        # Basic RTF conversion
        rtf_header = r"{\rtf1\ansi\deff0 {\fonttbl{\f0 Times New Roman;}}\f0\fs24 "
        rtf_footer = r"}"
        
        # Convert markdown to basic RTF
        rtf_content = markdown_content.replace('\n', r'\par ')
        rtf_content = rtf_content.replace('**', '')  # Remove bold markers
        rtf_content = rtf_content.replace('*', '')   # Remove italic markers
        
        return rtf_header + rtf_content + rtf_footer
    
    def save_output_file(self, output_path, content, format_type):
        """Save output file in the specified format"""
        if format_type in ['markdown', 'text', 'html', 'rtf']:
            # Text-based formats
            mode = 'w'
            encoding = 'utf-8'
            if format_type == 'rtf':
                encoding = 'latin-1'  # RTF typically uses latin-1
            
            with open(output_path, mode, encoding=encoding) as f:
                f.write(content)
                
        elif format_type == 'docx':
            # Create DOCX file
            if DOCX_AVAILABLE:
                doc = Document()
                
                # Parse markdown and add to document
                lines = content.split('\n') if content else []
                for line in lines:
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.strip():
                        doc.add_paragraph(line)
                
                doc.save(str(output_path))
            else:
                raise FileProcessingError("python-docx not installed")
                
        elif format_type == 'pdf':
            # Create PDF file
            if REPORTLAB_AVAILABLE:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                
                doc = SimpleDocTemplate(str(output_path), pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Parse markdown and add to PDF
                lines = content.split('\n') if content else []
                for line in lines:
                    if line.startswith('# '):
                        story.append(Paragraph(line[2:], styles['Title']))
                        story.append(Spacer(1, 12))
                    elif line.startswith('## '):
                        story.append(Paragraph(line[3:], styles['Heading1']))
                        story.append(Spacer(1, 12))
                    elif line.startswith('### '):
                        story.append(Paragraph(line[4:], styles['Heading2']))
                        story.append(Spacer(1, 12))
                    elif line.strip():
                        story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 12))
                
                doc.build(story)
            else:
                raise FileProcessingError("reportlab not installed")
    
    def convert_image_to_multiple_formats(self, file_path, output_format):
        """Convert image using OCR to specified format"""
        # First get text via OCR
        ocr_result = self.ocr_integration.process_file(
            str(file_path),
            language=self.ocr_language_var.get() if hasattr(self, 'ocr_language_var') else 'eng'
        )
        
        if ocr_result['success']:
            text_content = ocr_result['text']
            
            # Create appropriate header based on format
            if output_format == 'markdown':
                content = f"# {file_path.stem}\n\n"
                content += f"*Source: {file_path.name} (OCR)*\n\n"
                content += text_content
                if ocr_result.get('confidence'):
                    content += f"\n\n---\n*OCR Confidence: {ocr_result['confidence']:.1f}%*"
            elif output_format == 'text':
                content = f"{file_path.stem}\n"
                content += f"Source: {file_path.name} (OCR)\n\n"
                content += text_content
            elif output_format == 'html':
                content = f"<h1>{file_path.stem}</h1>\n"
                content += f"<p><em>Source: {file_path.name} (OCR)</em></p>\n"
                content += f"<p>{text_content.replace(chr(10), '<br>')}</p>"
            else:
                # For other formats, use text content directly
                content = text_content
            
            return content
        else:
            raise FileProcessingError(ocr_result.get('error', 'OCR failed'))

def main():
    """Application entry point"""
    try:
        # Try enhanced drag-and-drop support
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
    except ImportError:
        # Fall back to regular tkinter
        root = tk.Tk()
    
    # Create application
    app = BidirectionalDocumentConverter(root)
    
    # Center window
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    
    # Start application
    root.mainloop()

if __name__ == "__main__":
    main()
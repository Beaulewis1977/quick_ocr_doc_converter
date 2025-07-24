#!/usr/bin/env python3
"""
Advanced Document Converters for RTF, DOCX, and PDF output
Provides conversion functionality for complex document formats
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import tempfile
import subprocess
import sys

# RTF conversion
try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False

# DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# Alternative PDF generation
try:
    import fpdf
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

# CloudConvert integration
try:
    import cloudconvert
    HAS_CLOUDCONVERT = True
except ImportError:
    HAS_CLOUDCONVERT = False


class AdvancedConverter:
    """Advanced document converter for complex formats"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.logger = logging.getLogger(__name__)
        self.config = config or {}
        
        # Check CloudConvert API
        self.cloudconvert_client = None
        if HAS_CLOUDCONVERT and self.config.get('cloudconvert', {}).get('enabled'):
            api_key = self.config['cloudconvert'].get('api_key')
            if api_key:
                self.cloudconvert_client = cloudconvert.Api(api_key)
                self.logger.info("CloudConvert API initialized")
    
    def convert_to_rtf(self, input_file: str, output_file: str, content: Optional[str] = None) -> bool:
        """Convert text/markdown to RTF format"""
        try:
            # If content not provided, read from file
            if content is None:
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # Basic RTF conversion
            rtf_content = self._text_to_rtf(content)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(rtf_content)
            
            self.logger.info(f"Successfully converted to RTF: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"RTF conversion failed: {e}")
            return False
    
    def convert_from_rtf(self, input_file: str) -> Optional[str]:
        """Extract text from RTF file"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            
            if HAS_STRIPRTF:
                return rtf_to_text(rtf_content)
            else:
                # Basic RTF stripping
                return self._strip_rtf_basic(rtf_content)
                
        except Exception as e:
            self.logger.error(f"RTF extraction failed: {e}")
            return None
    
    def convert_to_docx(self, input_file: str, output_file: str, content: Optional[str] = None) -> bool:
        """Convert text/markdown to DOCX format"""
        try:
            if not HAS_PYTHON_DOCX:
                self.logger.error("python-docx not installed. Install with: pip install python-docx")
                return self._fallback_docx_conversion(input_file, output_file, content)
            
            # If content not provided, read from file
            if content is None:
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # Create DOCX document
            doc = Document()
            
            # Add title if markdown
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # Main title
                    doc.add_heading(line[2:], 0)
                elif line.startswith('## '):
                    # Subtitle
                    doc.add_heading(line[3:], 1)
                elif line.startswith('### '):
                    # Sub-subtitle
                    doc.add_heading(line[4:], 2)
                elif line.strip():
                    # Regular paragraph
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    # Empty line - add spacing
                    doc.add_paragraph()
            
            # Save document
            doc.save(output_file)
            self.logger.info(f"Successfully converted to DOCX: {output_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"DOCX conversion failed: {e}")
            return False
    
    def convert_to_pdf(self, input_file: str, output_file: str, content: Optional[str] = None) -> bool:
        """Convert text/markdown to PDF format"""
        try:
            # Try CloudConvert first if available
            if self.cloudconvert_client:
                return self._cloudconvert_to_pdf(input_file, output_file)
            
            # Use ReportLab if available
            if HAS_REPORTLAB:
                return self._reportlab_to_pdf(input_file, output_file, content)
            
            # Use FPDF as fallback
            if HAS_FPDF:
                return self._fpdf_to_pdf(input_file, output_file, content)
            
            # Last resort - use system tools
            return self._system_to_pdf(input_file, output_file)
            
        except Exception as e:
            self.logger.error(f"PDF conversion failed: {e}")
            return False
    
    def _text_to_rtf(self, text: str) -> str:
        """Convert plain text to RTF format"""
        # RTF header
        rtf = r'{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}'
        rtf += r'\f0\fs24 '  # Font size 12pt
        
        # Convert text to RTF
        # Replace special characters
        text = text.replace('\\', '\\\\')
        text = text.replace('{', '\\{')
        text = text.replace('}', '\\}')
        text = text.replace('\n', '\\par\n')
        
        rtf += text
        rtf += '}'
        
        return rtf
    
    def _strip_rtf_basic(self, rtf: str) -> str:
        """Basic RTF stripping without external library"""
        # Remove RTF control words and groups
        import re
        
        # Remove header
        rtf = re.sub(r'^\{\\rtf1.*?\\fs\d+\s*', '', rtf)
        
        # Remove control words
        rtf = re.sub(r'\\[a-z]+\d*\s*', ' ', rtf)
        
        # Remove groups
        rtf = re.sub(r'[{}]', '', rtf)
        
        # Handle special chars
        rtf = rtf.replace('\\par', '\n')
        rtf = rtf.replace('\\tab', '\t')
        
        return rtf.strip()
    
    def _reportlab_to_pdf(self, input_file: str, output_file: str, content: Optional[str] = None) -> bool:
        """Convert to PDF using ReportLab"""
        try:
            if content is None:
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # Create PDF
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Process content
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # Title
                    p = Paragraph(line[2:], styles['Title'])
                    story.append(p)
                    story.append(Spacer(1, 12))
                elif line.startswith('## '):
                    # Heading
                    p = Paragraph(line[3:], styles['Heading1'])
                    story.append(p)
                    story.append(Spacer(1, 6))
                elif line.strip():
                    # Normal text
                    p = Paragraph(line, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            self.logger.error(f"ReportLab PDF conversion failed: {e}")
            return False
    
    def _fpdf_to_pdf(self, input_file: str, output_file: str, content: Optional[str] = None) -> bool:
        """Convert to PDF using FPDF"""
        try:
            if content is None:
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # Create PDF
            pdf = fpdf.FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Process content
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # Title
                    pdf.set_font("Arial", 'B', 16)
                    pdf.cell(0, 10, txt=line[2:], ln=True, align='C')
                    pdf.set_font("Arial", size=12)
                elif line.startswith('## '):
                    # Heading
                    pdf.set_font("Arial", 'B', 14)
                    pdf.cell(0, 10, txt=line[3:], ln=True)
                    pdf.set_font("Arial", size=12)
                else:
                    # Normal text
                    pdf.multi_cell(0, 10, txt=line)
            
            # Save PDF
            pdf.output(output_file)
            return True
            
        except Exception as e:
            self.logger.error(f"FPDF conversion failed: {e}")
            return False
    
    def _system_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert to PDF using system tools"""
        try:
            # Try different system commands
            commands = [
                # LibreOffice/OpenOffice
                ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                 str(Path(output_file).parent), str(input_file)],
                
                # Pandoc
                ['pandoc', str(input_file), '-o', str(output_file)],
                
                # wkhtmltopdf (if HTML)
                ['wkhtmltopdf', str(input_file), str(output_file)],
            ]
            
            for cmd in commands:
                try:
                    result = subprocess.run(cmd, capture_output=True, text=True)
                    if result.returncode == 0:
                        # If LibreOffice was used, rename the output
                        if cmd[0] == 'soffice':
                            expected_output = Path(output_file).parent / f"{Path(input_file).stem}.pdf"
                            if expected_output.exists() and expected_output != Path(output_file):
                                expected_output.rename(output_file)
                        
                        self.logger.info(f"Successfully converted to PDF using {cmd[0]}")
                        return True
                except FileNotFoundError:
                    continue
            
            self.logger.error("No system PDF converter found")
            return False
            
        except Exception as e:
            self.logger.error(f"System PDF conversion failed: {e}")
            return False
    
    def _cloudconvert_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert to PDF using CloudConvert API"""
        try:
            # Create job
            job = self.cloudconvert_client.jobs.create({
                'tag': 'pdf_conversion',
                'tasks': {
                    'import-file': {
                        'operation': 'import/upload'
                    },
                    'convert-file': {
                        'operation': 'convert',
                        'input': 'import-file',
                        'output_format': 'pdf',
                        'engine': 'office'
                    },
                    'export-file': {
                        'operation': 'export/url',
                        'input': 'convert-file'
                    }
                }
            })
            
            # Upload file
            upload_task = job['tasks'][0]
            self.cloudconvert_client.tasks.upload(upload_task, input_file)
            
            # Wait for completion
            job = self.cloudconvert_client.jobs.wait(job['id'])
            
            # Download result
            export_task = None
            for task in job['tasks']:
                if task['operation'] == 'export/url':
                    export_task = task
                    break
            
            if export_task and export_task['status'] == 'finished':
                file_url = export_task['result']['files'][0]['url']
                
                # Download file
                import requests
                response = requests.get(file_url, verify=True)
                with open(output_file, 'wb') as f:
                    f.write(response.content)
                
                self.logger.info(f"Successfully converted to PDF using CloudConvert")
                return True
            
            return False
            
        except Exception as e:
            self.logger.error(f"CloudConvert PDF conversion failed: {e}")
            return False
    
    def _fallback_docx_conversion(self, input_file: str, output_file: str, content: Optional[str] = None) -> bool:
        """Fallback DOCX conversion using system tools"""
        try:
            # Try LibreOffice
            cmd = ['soffice', '--headless', '--convert-to', 'docx', '--outdir', 
                   str(Path(output_file).parent), str(input_file)]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                # Rename output if needed
                expected_output = Path(output_file).parent / f"{Path(input_file).stem}.docx"
                if expected_output.exists() and expected_output != Path(output_file):
                    expected_output.rename(output_file)
                return True
            
            return False
            
        except Exception:
            return False


# Global converter instance
advanced_converter = None

def get_advanced_converter(config: Optional[Dict[str, Any]] = None) -> AdvancedConverter:
    """Get or create advanced converter instance"""
    global advanced_converter
    if advanced_converter is None:
        advanced_converter = AdvancedConverter(config)
    return advanced_converter
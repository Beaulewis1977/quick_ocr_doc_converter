#!/usr/bin/env python3
"""
Universal Document Converter Complete - Enterprise Edition
Complete document and image conversion tool with OCR, threading, and advanced settings
Designed and built by Beau Lewis (blewisxx@gmail.com)

Features:
- Document conversion (DOCX, PDF, TXT, HTML, RTF, EPUB, ODT, XML, JSON, CSV)
- OCR functionality (JPG, PNG, TIFF, BMP, GIF, WebP, PDF with images)
- Multi-threaded processing with configurable workers
- Drag-and-drop support
- Advanced settings and configuration
- Cross-platform compatibility
- Professional GUI with responsive design
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
from pathlib import Path
import sys
import mimetypes
import re
import logging
import datetime
import json
from typing import Optional, Union, Dict, Any, List
import concurrent.futures
import time
import hashlib
from threading import Lock
import gc
import platform
import subprocess

# Optional dependencies
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

# Import OCR components
try:
    from ocr_engine.ocr_integration import OCRIntegration
    from ocr_engine.format_detector import OCRFormatDetector
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Custom Exception Classes
class DocumentConverterError(Exception):
    """Base exception class for document converter errors"""
    pass

class UnsupportedFormatError(DocumentConverterError):
    """Raised when an unsupported file format is encountered"""
    pass

class FileProcessingError(DocumentConverterError):
    """Raised when file processing fails"""
    pass

class ConfigurationManager:
    """Manages application configuration and settings"""
    
    def __init__(self):
        self.config_dir = Path.home() / '.document_converter'
        self.config_file = self.config_dir / 'settings.json'
        self.default_config = {
            'theme': 'default',
            'max_workers': 4,
            'output_format': 'markdown',
            'preserve_structure': True,
            'overwrite_existing': False,
            'ocr_enabled': False,
            'ocr_language': 'eng',
            'ocr_confidence_threshold': 60,
            'compression_enabled': False,
            'compression_level': 6,
            'logging_enabled': True,
            'log_level': 'INFO',
            'recent_folders': [],
            'window_geometry': None,
            'auto_open_output': True,
            'notification_enabled': True,
            'batch_size': 10,
            'memory_limit_mb': 500,
            'cache_enabled': True,
            'cache_size_mb': 100
        }
        self.config = self.load_config()
    
    def load_config(self) -> Dict[str, Any]:
        """Load configuration from file or create default"""
        if self.config_file.exists():
            try:
                with open(self.config_file, 'r') as f:
                    loaded_config = json.load(f)
                    # Merge with defaults to handle new keys
                    return {**self.default_config, **loaded_config}
            except Exception:
                return self.default_config.copy()
        return self.default_config.copy()
    
    def save_config(self):
        """Save current configuration to file"""
        try:
            self.config_dir.mkdir(parents=True, exist_ok=True)
            with open(self.config_file, 'w') as f:
                json.dump(self.config, f, indent=2)
        except Exception:
            pass
    
    def get(self, key: str, default=None):
        """Get configuration value"""
        return self.config.get(key, default)
    
    def set(self, key: str, value: Any):
        """Set configuration value"""
        self.config[key] = value
        self.save_config()

class DocumentConverterApp:
    """Main application class for the Universal Document Converter"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("Universal Document Converter Complete")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 700)
        
        # Initialize configuration
        self.config_manager = ConfigurationManager()
        
        # Initialize OCR if available
        self.ocr_integration = OCRIntegration() if OCR_AVAILABLE else None
        self.format_detector = OCRFormatDetector() if OCR_AVAILABLE else None
        
        # Processing state
        self.processing_thread = None
        self.cancel_processing = threading.Event()
        self.processing_lock = Lock()
        self.file_queue = []
        self.processed_files = set()
        
        # Statistics
        self.stats = {
            'total_files': 0,
            'processed': 0,
            'failed': 0,
            'skipped': 0,
            'start_time': None,
            'end_time': None
        }
        
        # GUI Variables
        self.setup_variables()
        
        # Apply saved window geometry
        if self.config_manager.get('window_geometry'):
            self.root.geometry(self.config_manager.get('window_geometry'))
        
        # Setup GUI
        self.setup_ui()
        
        # Bind events
        self.bind_events()
        
        # Load recent folders
        self.load_recent_folders()
        
        # Setup logging
        self.setup_logging()
        
        # Start memory monitor if available
        if PSUTIL_AVAILABLE:
            self.start_memory_monitor()
    
    def setup_variables(self):
        """Initialize GUI variables"""
        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.status_var = tk.StringVar(value="Ready")
        self.progress_var = tk.DoubleVar()
        self.memory_var = tk.StringVar(value="Memory: N/A")
        
        # Settings variables
        self.preserve_structure_var = tk.BooleanVar(value=self.config_manager.get('preserve_structure'))
        self.overwrite_var = tk.BooleanVar(value=self.config_manager.get('overwrite_existing'))
        self.ocr_enabled_var = tk.BooleanVar(value=self.config_manager.get('ocr_enabled'))
        self.compression_var = tk.BooleanVar(value=self.config_manager.get('compression_enabled'))
        self.max_workers_var = tk.IntVar(value=self.config_manager.get('max_workers'))
        self.auto_open_var = tk.BooleanVar(value=self.config_manager.get('auto_open_output'))
        self.notification_var = tk.BooleanVar(value=self.config_manager.get('notification_enabled'))
        
        # Set default output path
        desktop = Path.home() / "Desktop"
        default_output = desktop / "converted_documents"
        self.output_path.set(str(default_output))
    
    def setup_ui(self):
        """Setup the user interface"""
        # Configure styles
        self.setup_styles()
        
        # Create main container
        main_container = ttk.Frame(self.root)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create menu bar
        self.create_menu_bar()
        
        # Create toolbar
        self.create_toolbar(main_container)
        
        # Create main content area
        content_frame = ttk.Frame(main_container)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Left panel - Input/Output and Settings
        left_panel = ttk.Frame(content_frame)
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        # Input/Output section
        self.create_io_section(left_panel)
        
        # Settings section
        self.create_settings_section(left_panel)
        
        # Right panel - File list and progress
        right_panel = ttk.Frame(content_frame)
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        # File list section
        self.create_file_list_section(right_panel)
        
        # Progress section
        self.create_progress_section(right_panel)
        
        # Status bar
        self.create_status_bar(main_container)
        
        # Configure drag and drop
        self.setup_drag_drop()
    
    def setup_styles(self):
        """Configure ttk styles"""
        style = ttk.Style()
        
        # Configure button styles
        style.configure('Accent.TButton', font=('Arial', 10, 'bold'))
        style.configure('Success.TButton', foreground='green')
        style.configure('Danger.TButton', foreground='red')
        
        # Configure frame styles
        style.configure('Card.TFrame', relief='solid', borderwidth=1)
    
    def create_menu_bar(self):
        """Create application menu bar"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Select Input Folder", command=self.browse_input_folder)
        file_menu.add_command(label="Select Output Folder", command=self.browse_output_folder)
        file_menu.add_separator()
        file_menu.add_command(label="Recent Folders", command=self.show_recent_folders)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.on_closing)
        
        # Tools menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Tools", menu=tools_menu)
        tools_menu.add_command(label="Batch Convert", command=self.batch_convert)
        tools_menu.add_command(label="OCR Settings", command=self.show_ocr_settings)
        tools_menu.add_separator()
        tools_menu.add_command(label="Clear Cache", command=self.clear_cache)
        tools_menu.add_command(label="View Logs", command=self.view_logs)
        
        # Settings menu
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Settings", menu=settings_menu)
        settings_menu.add_command(label="Preferences", command=self.show_preferences)
        settings_menu.add_command(label="Advanced Settings", command=self.show_advanced_settings)
        
        # Help menu
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="Documentation", command=self.show_documentation)
        help_menu.add_command(label="Keyboard Shortcuts", command=self.show_shortcuts)
        help_menu.add_separator()
        help_menu.add_command(label="About", command=self.show_about)
    
    def create_toolbar(self, parent):
        """Create application toolbar"""
        toolbar = ttk.Frame(parent)
        toolbar.pack(fill=tk.X, padx=10, pady=5)
        
        # Toolbar buttons
        ttk.Button(toolbar, text="📁 Input", command=self.browse_input_folder).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="📂 Output", command=self.browse_output_folder).pack(side=tk.LEFT, padx=2)
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        ttk.Button(toolbar, text="▶️ Convert", command=self.start_conversion, 
                  style='Accent.TButton').pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="⏸️ Pause", command=self.pause_conversion).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="⏹️ Stop", command=self.stop_conversion,
                  style='Danger.TButton').pack(side=tk.LEFT, padx=2)
        
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
        
        # Thread selector
        ttk.Label(toolbar, text="Threads:").pack(side=tk.LEFT, padx=(5, 2))
        thread_spinbox = ttk.Spinbox(toolbar, from_=1, to=16, width=5,
                                    textvariable=self.max_workers_var)
        thread_spinbox.pack(side=tk.LEFT, padx=2)
        
        # Memory usage
        if PSUTIL_AVAILABLE:
            ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, padx=5, fill=tk.Y)
            ttk.Label(toolbar, textvariable=self.memory_var).pack(side=tk.LEFT, padx=5)
    
    def create_io_section(self, parent):
        """Create input/output selection section"""
        io_frame = ttk.LabelFrame(parent, text="Input/Output", padding=10)
        io_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Input folder
        ttk.Label(io_frame, text="Input Folder:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(io_frame, textvariable=self.input_path, width=50).grid(row=0, column=1, padx=5, sticky=tk.EW)
        ttk.Button(io_frame, text="Browse", command=self.browse_input_folder).grid(row=0, column=2)
        
        # Output folder
        ttk.Label(io_frame, text="Output Folder:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(io_frame, textvariable=self.output_path, width=50).grid(row=1, column=1, padx=5, sticky=tk.EW)
        ttk.Button(io_frame, text="Browse", command=self.browse_output_folder).grid(row=1, column=2)
        
        io_frame.columnconfigure(1, weight=1)
    
    def create_settings_section(self, parent):
        """Create settings section"""
        settings_frame = ttk.LabelFrame(parent, text="Settings", padding=10)
        settings_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create notebook for organized settings
        notebook = ttk.Notebook(settings_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # General settings tab
        general_tab = ttk.Frame(notebook)
        notebook.add(general_tab, text="General")
        
        ttk.Checkbutton(general_tab, text="Preserve folder structure",
                       variable=self.preserve_structure_var).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(general_tab, text="Overwrite existing files",
                       variable=self.overwrite_var).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(general_tab, text="Auto-open output folder",
                       variable=self.auto_open_var).pack(anchor=tk.W, pady=2)
        ttk.Checkbutton(general_tab, text="Show notifications",
                       variable=self.notification_var).pack(anchor=tk.W, pady=2)
        
        # OCR settings tab
        if OCR_AVAILABLE:
            ocr_tab = ttk.Frame(notebook)
            notebook.add(ocr_tab, text="OCR")
            
            ttk.Checkbutton(ocr_tab, text="Enable OCR for images",
                           variable=self.ocr_enabled_var,
                           command=self.toggle_ocr).pack(anchor=tk.W, pady=2)
            
            # OCR status
            self.ocr_status_label = ttk.Label(ocr_tab, text="")
            self.ocr_status_label.pack(anchor=tk.W, pady=5)
            
            # OCR language selector
            lang_frame = ttk.Frame(ocr_tab)
            lang_frame.pack(anchor=tk.W, pady=5)
            ttk.Label(lang_frame, text="Language:").pack(side=tk.LEFT, padx=(0, 5))
            self.ocr_language_var = tk.StringVar(value="eng")
            lang_combo = ttk.Combobox(lang_frame, textvariable=self.ocr_language_var,
                                     values=["eng", "fra", "deu", "spa", "ita", "por", "rus", "jpn", "chi_sim"],
                                     width=15)
            lang_combo.pack(side=tk.LEFT)
            
            # Check OCR status
            self.check_ocr_status()
        
        # Advanced settings tab
        advanced_tab = ttk.Frame(notebook)
        notebook.add(advanced_tab, text="Advanced")
        
        ttk.Checkbutton(advanced_tab, text="Enable compression",
                       variable=self.compression_var).pack(anchor=tk.W, pady=2)
        
        # Batch size setting
        batch_frame = ttk.Frame(advanced_tab)
        batch_frame.pack(anchor=tk.W, pady=5)
        ttk.Label(batch_frame, text="Batch size:").pack(side=tk.LEFT, padx=(0, 5))
        self.batch_size_var = tk.IntVar(value=self.config_manager.get('batch_size'))
        ttk.Spinbox(batch_frame, from_=1, to=100, width=10,
                   textvariable=self.batch_size_var).pack(side=tk.LEFT)
        
        # Memory limit setting
        memory_frame = ttk.Frame(advanced_tab)
        memory_frame.pack(anchor=tk.W, pady=5)
        ttk.Label(memory_frame, text="Memory limit (MB):").pack(side=tk.LEFT, padx=(0, 5))
        self.memory_limit_var = tk.IntVar(value=self.config_manager.get('memory_limit_mb'))
        ttk.Spinbox(memory_frame, from_=100, to=4000, width=10,
                   textvariable=self.memory_limit_var).pack(side=tk.LEFT)
    
    def create_file_list_section(self, parent):
        """Create file list section"""
        list_frame = ttk.LabelFrame(parent, text="Files to Convert", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Create treeview for file list
        columns = ('Size', 'Type', 'Status')
        self.file_tree = ttk.Treeview(list_frame, columns=columns, show='tree headings', height=15)
        
        # Configure columns
        self.file_tree.heading('#0', text='File')
        self.file_tree.heading('Size', text='Size')
        self.file_tree.heading('Type', text='Type')
        self.file_tree.heading('Status', text='Status')
        
        self.file_tree.column('#0', width=300)
        self.file_tree.column('Size', width=80)
        self.file_tree.column('Type', width=80)
        self.file_tree.column('Status', width=100)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        hsb = ttk.Scrollbar(list_frame, orient=tk.HORIZONTAL, command=self.file_tree.xview)
        self.file_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Pack widgets
        self.file_tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        
        list_frame.rowconfigure(0, weight=1)
        list_frame.columnconfigure(0, weight=1)
        
        # Context menu
        self.create_file_context_menu()
    
    def create_progress_section(self, parent):
        """Create progress section"""
        progress_frame = ttk.LabelFrame(parent, text="Progress", padding=10)
        progress_frame.pack(fill=tk.X)
        
        # Overall progress
        ttk.Label(progress_frame, text="Overall Progress:").pack(anchor=tk.W)
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var,
                                           maximum=100, length=400)
        self.progress_bar.pack(fill=tk.X, pady=5)
        
        # Statistics
        stats_frame = ttk.Frame(progress_frame)
        stats_frame.pack(fill=tk.X)
        
        self.stats_labels = {}
        stats = ['Total Files', 'Processed', 'Failed', 'Skipped', 'Time Elapsed']
        for i, stat in enumerate(stats):
            label = ttk.Label(stats_frame, text=f"{stat}: 0")
            label.grid(row=i//3, column=(i%3)*2, sticky=tk.W, padx=5)
            self.stats_labels[stat.lower().replace(' ', '_')] = label
    
    def create_status_bar(self, parent):
        """Create status bar"""
        status_frame = ttk.Frame(parent)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        
        ttk.Label(status_frame, textvariable=self.status_var).pack(side=tk.LEFT, padx=10)
        
        # Add grip for window resizing
        ttk.Sizegrip(status_frame).pack(side=tk.RIGHT)
    
    def create_file_context_menu(self):
        """Create context menu for file list"""
        self.file_context_menu = tk.Menu(self.root, tearoff=0)
        self.file_context_menu.add_command(label="Remove", command=self.remove_selected_files)
        self.file_context_menu.add_command(label="Clear All", command=self.clear_file_list)
        self.file_context_menu.add_separator()
        self.file_context_menu.add_command(label="Show in Explorer", command=self.show_in_explorer)
        
        self.file_tree.bind('<Button-3>', self.show_file_context_menu)
    
    def show_file_context_menu(self, event):
        """Show context menu for file list"""
        try:
            self.file_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.file_context_menu.grab_release()
    
    def bind_events(self):
        """Bind keyboard and window events"""
        # Keyboard shortcuts
        self.root.bind('<Control-o>', lambda e: self.browse_input_folder())
        self.root.bind('<Control-s>', lambda e: self.browse_output_folder())
        self.root.bind('<Control-r>', lambda e: self.start_conversion())
        self.root.bind('<Escape>', lambda e: self.stop_conversion())
        self.root.bind('<F1>', lambda e: self.show_documentation())
        
        # Window events
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.root.bind('<Configure>', self.on_window_configure)
    
    def setup_drag_drop(self):
        """Setup drag and drop functionality"""
        try:
            from tkinterdnd2 import TkinterDnD, DND_FILES
            # Make the root window a drop target
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self.on_drop)
        except ImportError:
            # tkinterdnd2 not available
            pass
    
    def on_drop(self, event):
        """Handle drag and drop events"""
        files = self.root.tk.splitlist(event.data)
        if files:
            # Use the first dropped item
            dropped_path = files[0]
            if os.path.isdir(dropped_path):
                self.input_path.set(dropped_path)
                self.scan_input_folder()
            else:
                # If it's a file, use its parent directory
                parent_dir = os.path.dirname(dropped_path)
                self.input_path.set(parent_dir)
                self.scan_input_folder()
    
    def setup_logging(self):
        """Setup application logging"""
        if self.config_manager.get('logging_enabled'):
            log_dir = Path.home() / '.document_converter' / 'logs'
            log_dir.mkdir(parents=True, exist_ok=True)
            
            log_file = log_dir / f"converter_{datetime.datetime.now().strftime('%Y%m%d')}.log"
            
            logging.basicConfig(
                level=getattr(logging, self.config_manager.get('log_level')),
                format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                handlers=[
                    logging.FileHandler(log_file),
                    logging.StreamHandler()
                ]
            )
            self.logger = logging.getLogger('DocumentConverter')
        else:
            self.logger = None
    
    def start_memory_monitor(self):
        """Start memory usage monitoring"""
        def update_memory():
            if PSUTIL_AVAILABLE:
                process = psutil.Process()
                memory_mb = process.memory_info().rss / 1024 / 1024
                self.memory_var.set(f"Memory: {memory_mb:.1f} MB")
            self.root.after(2000, update_memory)
        
        update_memory()
    
    def browse_input_folder(self):
        """Browse for input folder"""
        folder = filedialog.askdirectory(title="Select input folder")
        if folder:
            self.input_path.set(folder)
            self.add_to_recent_folders(folder)
            self.scan_input_folder()
    
    def browse_output_folder(self):
        """Browse for output folder"""
        folder = filedialog.askdirectory(title="Select output folder")
        if folder:
            self.output_path.set(folder)
            self.add_to_recent_folders(folder)
    
    def scan_input_folder(self):
        """Scan input folder for convertible files"""
        if not self.input_path.get():
            return
        
        self.clear_file_list()
        self.status_var.set("Scanning folder...")
        
        # Define supported formats
        doc_patterns = ['*.docx', '*.pdf', '*.txt', '*.rtf', '*.odt', '*.html', '*.htm', 
                       '*.epub', '*.xml', '*.json', '*.csv']
        
        # Add image patterns if OCR is enabled
        if self.ocr_enabled_var.get() and OCR_AVAILABLE:
            image_patterns = ['*.jpg', '*.jpeg', '*.png', '*.tiff', '*.tif', '*.bmp', 
                            '*.gif', '*.webp']
            doc_patterns.extend(image_patterns)
        
        input_path = Path(self.input_path.get())
        files_found = []
        
        for pattern in doc_patterns:
            files_found.extend(input_path.rglob(pattern))
        
        # Filter out temporary files
        files_found = [f for f in files_found if not f.name.startswith('~$')]
        
        # Add files to tree
        for file_path in files_found:
            size = self.format_file_size(file_path.stat().st_size)
            file_type = file_path.suffix.upper()[1:]
            rel_path = file_path.relative_to(input_path)
            
            self.file_tree.insert('', 'end', text=str(rel_path), 
                                 values=(size, file_type, 'Pending'))
        
        self.stats['total_files'] = len(files_found)
        self.update_stats_display()
        self.status_var.set(f"Found {len(files_found)} files to convert")
    
    def format_file_size(self, size_bytes):
        """Format file size for display"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
    
    def clear_file_list(self):
        """Clear the file list"""
        self.file_tree.delete(*self.file_tree.get_children())
        self.stats = {
            'total_files': 0,
            'processed': 0,
            'failed': 0,
            'skipped': 0,
            'start_time': None,
            'end_time': None
        }
        self.update_stats_display()
    
    def remove_selected_files(self):
        """Remove selected files from the list"""
        selected = self.file_tree.selection()
        for item in selected:
            self.file_tree.delete(item)
        self.stats['total_files'] = len(self.file_tree.get_children())
        self.update_stats_display()
    
    def show_in_explorer(self):
        """Show selected file in file explorer"""
        selected = self.file_tree.selection()
        if selected:
            file_path = self.file_tree.item(selected[0])['text']
            full_path = Path(self.input_path.get()) / file_path
            
            if platform.system() == 'Windows':
                subprocess.run(['explorer', '/select,', str(full_path)])
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', '-R', str(full_path)])
            else:  # Linux
                subprocess.run(['xdg-open', str(full_path.parent)])
    
    def toggle_ocr(self):
        """Toggle OCR mode"""
        if self.ocr_enabled_var.get():
            self.scan_input_folder()  # Rescan to include images
            self.status_var.set("OCR mode enabled")
        else:
            self.scan_input_folder()  # Rescan to exclude images
            self.status_var.set("OCR mode disabled")
    
    def check_ocr_status(self):
        """Check OCR availability and update status"""
        if not OCR_AVAILABLE:
            self.ocr_status_label.config(text="OCR not available - install dependencies", 
                                       foreground='red')
            return
        
        try:
            availability = self.ocr_integration.check_availability()
            if availability['available']:
                self.ocr_status_label.config(text=f"✓ {availability['message']}", 
                                           foreground='green')
            else:
                self.ocr_status_label.config(text=f"⚠ {availability['message']}", 
                                           foreground='orange')
        except Exception as e:
            self.ocr_status_label.config(text=f"Error: {str(e)}", 
                                       foreground='red')
    
    def start_conversion(self):
        """Start the conversion process"""
        if self.processing_thread and self.processing_thread.is_alive():
            messagebox.showwarning("Processing", "Conversion already in progress!")
            return
        
        if not self.validate_inputs():
            return
        
        # Reset stats
        self.stats = {
            'total_files': len(self.file_tree.get_children()),
            'processed': 0,
            'failed': 0,
            'skipped': 0,
            'start_time': time.time(),
            'end_time': None
        }
        
        # Clear cancel flag
        self.cancel_processing.clear()
        
        # Update UI
        self.status_var.set("Starting conversion...")
        self.progress_var.set(0)
        
        # Save current settings
        self.save_current_settings()
        
        # Start processing thread
        self.processing_thread = threading.Thread(target=self.process_files)
        self.processing_thread.daemon = True
        self.processing_thread.start()
    
    def validate_inputs(self):
        """Validate user inputs"""
        if not self.input_path.get():
            messagebox.showerror("Error", "Please select an input folder")
            return False
        
        if not os.path.exists(self.input_path.get()):
            messagebox.showerror("Error", "Input folder does not exist")
            return False
        
        if not self.output_path.get():
            messagebox.showerror("Error", "Please select an output folder")
            return False
        
        if len(self.file_tree.get_children()) == 0:
            messagebox.showerror("Error", "No files to convert")
            return False
        
        return True
    
    def save_current_settings(self):
        """Save current settings to config"""
        self.config_manager.set('preserve_structure', self.preserve_structure_var.get())
        self.config_manager.set('overwrite_existing', self.overwrite_var.get())
        self.config_manager.set('ocr_enabled', self.ocr_enabled_var.get())
        self.config_manager.set('compression_enabled', self.compression_var.get())
        self.config_manager.set('max_workers', self.max_workers_var.get())
        self.config_manager.set('auto_open_output', self.auto_open_var.get())
        self.config_manager.set('notification_enabled', self.notification_var.get())
        self.config_manager.set('batch_size', self.batch_size_var.get())
        self.config_manager.set('memory_limit_mb', self.memory_limit_var.get())
        if hasattr(self, 'ocr_language_var'):
            self.config_manager.set('ocr_language', self.ocr_language_var.get())
    
    def process_files(self):
        """Process files in a separate thread"""
        try:
            input_dir = Path(self.input_path.get())
            output_dir = Path(self.output_path.get())
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Get all files from tree
            files_to_process = []
            for item in self.file_tree.get_children():
                file_path = input_dir / self.file_tree.item(item)['text']
                files_to_process.append((item, file_path))
            
            # Process files in batches with thread pool
            max_workers = self.max_workers_var.get()
            batch_size = self.batch_size_var.get()
            
            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Process in batches
                for i in range(0, len(files_to_process), batch_size):
                    if self.cancel_processing.is_set():
                        break
                    
                    batch = files_to_process[i:i+batch_size]
                    futures = []
                    
                    for item_id, file_path in batch:
                        if self.cancel_processing.is_set():
                            break
                        
                        # Update status
                        self.update_file_status(item_id, "Processing...")
                        
                        # Submit conversion task
                        future = executor.submit(self.convert_single_file, 
                                               item_id, file_path, input_dir, output_dir)
                        futures.append(future)
                    
                    # Wait for batch to complete
                    for future in concurrent.futures.as_completed(futures):
                        if self.cancel_processing.is_set():
                            break
                        
                        try:
                            result = future.result()
                            # Result handled in convert_single_file
                        except Exception as e:
                            if self.logger:
                                self.logger.error(f"Conversion error: {str(e)}")
                    
                    # Force garbage collection after each batch
                    gc.collect()
                    
                    # Check memory usage
                    if PSUTIL_AVAILABLE:
                        process = psutil.Process()
                        memory_mb = process.memory_info().rss / 1024 / 1024
                        if memory_mb > self.memory_limit_var.get():
                            self.status_var.set("Pausing for memory cleanup...")
                            time.sleep(2)
                            gc.collect()
            
            # Conversion complete
            self.stats['end_time'] = time.time()
            self.show_completion_summary()
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"Process files error: {str(e)}")
            self.status_var.set(f"Error: {str(e)}")
        finally:
            self.processing_thread = None
    
    def convert_single_file(self, item_id, file_path, input_dir, output_dir):
        """Convert a single file"""
        try:
            # Calculate output path
            if self.preserve_structure_var.get():
                rel_path = file_path.relative_to(input_dir)
                output_path = output_dir / rel_path.with_suffix('.md')
            else:
                output_path = output_dir / f"{file_path.stem}.md"
            
            # Create output directory if needed
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Check if file exists and skip if needed
            if output_path.exists() and not self.overwrite_var.get():
                self.update_file_status(item_id, "Skipped")
                self.stats['skipped'] += 1
                self.update_progress()
                return
            
            # Convert based on file type
            file_ext = file_path.suffix.lower()
            content = None
            
            # Check if it's an image file for OCR
            image_extensions = ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif', '.webp']
            
            if file_ext in image_extensions and self.ocr_enabled_var.get() and OCR_AVAILABLE:
                content = self.convert_image_to_markdown(file_path)
            elif file_ext == '.docx':
                content = self.convert_docx_to_markdown(file_path)
            elif file_ext == '.pdf':
                # Check if PDF is image-based
                if self.ocr_enabled_var.get() and OCR_AVAILABLE and self.is_pdf_image_based(file_path):
                    content = self.convert_image_to_markdown(file_path)
                else:
                    content = self.convert_pdf_to_markdown(file_path)
            elif file_ext in ['.txt', '.md']:
                content = self.convert_txt_to_markdown(file_path)
            elif file_ext in ['.html', '.htm']:
                content = self.convert_html_to_markdown(file_path)
            elif file_ext == '.rtf':
                content = self.convert_rtf_to_markdown(file_path)
            elif file_ext == '.odt':
                content = self.convert_odt_to_markdown(file_path)
            elif file_ext == '.epub':
                content = self.convert_epub_to_markdown(file_path)
            elif file_ext == '.json':
                content = self.convert_json_to_markdown(file_path)
            elif file_ext == '.csv':
                content = self.convert_csv_to_markdown(file_path)
            elif file_ext == '.xml':
                content = self.convert_xml_to_markdown(file_path)
            else:
                raise UnsupportedFormatError(f"Unsupported format: {file_ext}")
            
            # Save converted content
            if content:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # Apply compression if enabled
                if self.compression_var.get():
                    self.compress_file(output_path)
                
                self.update_file_status(item_id, "✓ Complete")
                self.stats['processed'] += 1
            else:
                raise FileProcessingError("No content extracted")
            
        except Exception as e:
            self.update_file_status(item_id, f"✗ Failed: {str(e)}")
            self.stats['failed'] += 1
            if self.logger:
                self.logger.error(f"Failed to convert {file_path}: {str(e)}")
        
        finally:
            self.update_progress()
    
    def convert_docx_to_markdown(self, file_path):
        """Convert DOCX to Markdown"""
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            content = [f"# {file_path.stem}\n"]
            
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    content.append(f"{'#' * level} {paragraph.text}\n")
                elif paragraph.text.strip():
                    content.append(f"{paragraph.text}\n")
            
            # Handle tables
            for table in doc.tables:
                content.append("\n")
                for i, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    content.append(f"| {row_text} |")
                    if i == 0:
                        content.append("|" + " --- |" * len(row.cells))
                content.append("\n")
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"DOCX conversion failed: {str(e)}")
    
    def convert_pdf_to_markdown(self, file_path):
        """Convert PDF to Markdown"""
        try:
            import PyPDF2
            
            content = [f"# {file_path.stem}\n"]
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    content.append(f"\n## Page {page_num}\n")
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"PDF conversion failed: {str(e)}")
    
    def convert_txt_to_markdown(self, file_path):
        """Convert TXT to Markdown"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return f"# {file_path.stem}\n\n{content}"
                except UnicodeDecodeError:
                    continue
            
            raise FileProcessingError(f"Could not decode file with any encoding")
            
        except Exception as e:
            raise FileProcessingError(f"TXT conversion failed: {str(e)}")
    
    def convert_html_to_markdown(self, file_path):
        """Convert HTML to Markdown"""
        try:
            from bs4 import BeautifulSoup
            import html2text
            
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Use html2text if available
            try:
                h = html2text.HTML2Text()
                h.ignore_links = False
                return h.handle(html_content)
            except:
                # Fallback to BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                return f"# {file_path.stem}\n\n{soup.get_text()}"
                
        except Exception as e:
            raise FileProcessingError(f"HTML conversion failed: {str(e)}")
    
    def convert_rtf_to_markdown(self, file_path):
        """Convert RTF to Markdown"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
            
            text = rtf_to_text(rtf_content)
            return f"# {file_path.stem}\n\n{text}"
            
        except Exception as e:
            raise FileProcessingError(f"RTF conversion failed: {str(e)}")
    
    def convert_odt_to_markdown(self, file_path):
        """Convert ODT to Markdown"""
        try:
            from odf import text, teletype
            from odf.opendocument import load
            
            doc = load(str(file_path))
            content = [f"# {file_path.stem}\n"]
            
            for element in doc.getElementsByType(text.P):
                content.append(teletype.extractText(element))
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"ODT conversion failed: {str(e)}")
    
    def convert_epub_to_markdown(self, file_path):
        """Convert EPUB to Markdown"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(str(file_path))
            content = [f"# {file_path.stem}\n"]
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    content.append(soup.get_text())
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"EPUB conversion failed: {str(e)}")
    
    def convert_json_to_markdown(self, file_path):
        """Convert JSON to Markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            content = [f"# {file_path.stem}\n"]
            content.append("```json")
            content.append(json.dumps(data, indent=2))
            content.append("```")
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"JSON conversion failed: {str(e)}")
    
    def convert_csv_to_markdown(self, file_path):
        """Convert CSV to Markdown table"""
        try:
            import csv
            
            content = [f"# {file_path.stem}\n"]
            
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                rows = list(csv_reader)
                
                if rows:
                    # Header
                    content.append("| " + " | ".join(rows[0]) + " |")
                    content.append("|" + " --- |" * len(rows[0]))
                    
                    # Data rows
                    for row in rows[1:]:
                        content.append("| " + " | ".join(row) + " |")
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"CSV conversion failed: {str(e)}")
    
    def convert_xml_to_markdown(self, file_path):
        """Convert XML to Markdown"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(str(file_path))
            root = tree.getroot()
            
            content = [f"# {file_path.stem}\n"]
            content.append("```xml")
            content.append(ET.tostring(root, encoding='unicode'))
            content.append("```")
            
            return "\n".join(content)
            
        except Exception as e:
            raise FileProcessingError(f"XML conversion failed: {str(e)}")
    
    def convert_image_to_markdown(self, file_path):
        """Convert image to Markdown using OCR"""
        try:
            if not OCR_AVAILABLE:
                raise FileProcessingError("OCR not available")
            
            result = self.ocr_integration.process_file(str(file_path), 
                                                     language=self.ocr_language_var.get())
            
            if result['success']:
                content = f"# {file_path.stem}\n\n"
                content += f"*Source: {file_path.name} (OCR)*\n\n"
                content += result['text']
                
                if result.get('confidence'):
                    content += f"\n\n---\n*OCR Confidence: {result['confidence']:.1f}%*"
                
                return content
            else:
                raise FileProcessingError(result.get('error', 'OCR failed'))
                
        except Exception as e:
            raise FileProcessingError(f"Image conversion failed: {str(e)}")
    
    def is_pdf_image_based(self, file_path):
        """Check if PDF is image-based (scanned)"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check first few pages
                pages_to_check = min(3, len(pdf_reader.pages))
                total_text_length = 0
                
                for i in range(pages_to_check):
                    page = pdf_reader.pages[i]
                    text = page.extract_text()
                    total_text_length += len(text.strip())
                
                # If very little text, likely image-based
                return total_text_length < 100
                
        except Exception:
            return False
    
    def compress_file(self, file_path):
        """Compress the output file"""
        try:
            import gzip
            
            with open(file_path, 'rb') as f_in:
                with gzip.open(f"{file_path}.gz", 'wb', compresslevel=6) as f_out:
                    f_out.writelines(f_in)
            
            # Remove original file
            os.remove(file_path)
            
        except Exception:
            pass  # Compression is optional
    
    def update_file_status(self, item_id, status):
        """Update file status in tree"""
        def update():
            try:
                values = list(self.file_tree.item(item_id)['values'])
                values[2] = status
                self.file_tree.item(item_id, values=values)
            except Exception:
                pass
        
        self.root.after(0, update)
    
    def update_progress(self):
        """Update progress bar and statistics"""
        def update():
            total = self.stats['total_files']
            completed = self.stats['processed'] + self.stats['failed'] + self.stats['skipped']
            
            if total > 0:
                progress = (completed / total) * 100
                self.progress_var.set(progress)
            
            self.update_stats_display()
            
            # Update status
            self.status_var.set(f"Processing: {completed}/{total} files")
        
        self.root.after(0, update)
    
    def update_stats_display(self):
        """Update statistics display"""
        self.stats_labels['total_files'].config(text=f"Total Files: {self.stats['total_files']}")
        self.stats_labels['processed'].config(text=f"Processed: {self.stats['processed']}")
        self.stats_labels['failed'].config(text=f"Failed: {self.stats['failed']}")
        self.stats_labels['skipped'].config(text=f"Skipped: {self.stats['skipped']}")
        
        # Calculate elapsed time
        if self.stats['start_time']:
            elapsed = time.time() - self.stats['start_time']
            elapsed_str = time.strftime('%H:%M:%S', time.gmtime(elapsed))
            self.stats_labels['time_elapsed'].config(text=f"Time Elapsed: {elapsed_str}")
    
    def show_completion_summary(self):
        """Show conversion completion summary"""
        elapsed = self.stats['end_time'] - self.stats['start_time']
        elapsed_str = time.strftime('%H:%M:%S', time.gmtime(elapsed))
        
        summary = f"Conversion Complete!\n\n"
        summary += f"Total Files: {self.stats['total_files']}\n"
        summary += f"Processed: {self.stats['processed']}\n"
        summary += f"Failed: {self.stats['failed']}\n"
        summary += f"Skipped: {self.stats['skipped']}\n"
        summary += f"Time Elapsed: {elapsed_str}\n\n"
        summary += f"Output saved to: {self.output_path.get()}"
        
        self.status_var.set("Conversion complete!")
        
        if self.notification_var.get():
            messagebox.showinfo("Conversion Complete", summary)
        
        if self.auto_open_var.get():
            self.open_output_folder()
        
        if self.logger:
            self.logger.info(f"Conversion complete: {self.stats}")
    
    def pause_conversion(self):
        """Pause the conversion process"""
        # Not implemented in this version
        messagebox.showinfo("Info", "Pause functionality not yet implemented")
    
    def stop_conversion(self):
        """Stop the conversion process"""
        if self.processing_thread and self.processing_thread.is_alive():
            self.cancel_processing.set()
            self.status_var.set("Stopping conversion...")
            messagebox.showinfo("Stopping", "Conversion will stop after current file")
    
    def open_output_folder(self):
        """Open the output folder in file explorer"""
        output_path = self.output_path.get()
        if output_path and os.path.exists(output_path):
            if platform.system() == 'Windows':
                os.startfile(output_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', output_path])
            else:  # Linux
                subprocess.run(['xdg-open', output_path])
    
    def batch_convert(self):
        """Show batch conversion dialog"""
        messagebox.showinfo("Batch Convert", 
                          "Batch conversion allows processing multiple folders.\n"
                          "This feature is coming soon!")
    
    def show_ocr_settings(self):
        """Show OCR settings dialog"""
        if not OCR_AVAILABLE:
            messagebox.showwarning("OCR Not Available", 
                                 "OCR functionality is not available.\n"
                                 "Please install OCR dependencies.")
            return
        
        # Create OCR settings dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("OCR Settings")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Add OCR settings controls
        ttk.Label(dialog, text="OCR Settings", font=('Arial', 12, 'bold')).pack(pady=10)
        
        # Language selection
        lang_frame = ttk.Frame(dialog)
        lang_frame.pack(pady=10)
        ttk.Label(lang_frame, text="Language:").pack(side=tk.LEFT, padx=5)
        lang_combo = ttk.Combobox(lang_frame, textvariable=self.ocr_language_var,
                                 values=["eng", "fra", "deu", "spa", "ita", "por", "rus", "jpn", "chi_sim"],
                                 width=20)
        lang_combo.pack(side=tk.LEFT)
        
        # Confidence threshold
        threshold_frame = ttk.Frame(dialog)
        threshold_frame.pack(pady=10)
        ttk.Label(threshold_frame, text="Confidence Threshold:").pack(side=tk.LEFT, padx=5)
        self.ocr_threshold_var = tk.IntVar(value=self.config_manager.get('ocr_confidence_threshold'))
        ttk.Scale(threshold_frame, from_=0, to=100, variable=self.ocr_threshold_var,
                 orient=tk.HORIZONTAL, length=200).pack(side=tk.LEFT)
        ttk.Label(threshold_frame, textvariable=self.ocr_threshold_var).pack(side=tk.LEFT, padx=5)
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=20)
        ttk.Button(button_frame, text="Save", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def clear_cache(self):
        """Clear application cache"""
        if OCR_AVAILABLE and self.ocr_integration:
            try:
                self.ocr_integration.clear_cache()
                messagebox.showinfo("Cache Cleared", "OCR cache has been cleared successfully.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to clear cache: {str(e)}")
        else:
            messagebox.showinfo("Info", "No cache to clear.")
    
    def view_logs(self):
        """View application logs"""
        log_dir = Path.home() / '.document_converter' / 'logs'
        if log_dir.exists():
            if platform.system() == 'Windows':
                os.startfile(str(log_dir))
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', str(log_dir)])
            else:  # Linux
                subprocess.run(['xdg-open', str(log_dir)])
        else:
            messagebox.showinfo("No Logs", "No log files found.")
    
    def show_preferences(self):
        """Show preferences dialog"""
        # Create preferences dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Preferences")
        dialog.geometry("500x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create notebook for organized preferences
        notebook = ttk.Notebook(dialog)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # General tab
        general_tab = ttk.Frame(notebook)
        notebook.add(general_tab, text="General")
        
        # Appearance tab
        appearance_tab = ttk.Frame(notebook)
        notebook.add(appearance_tab, text="Appearance")
        
        # Performance tab
        performance_tab = ttk.Frame(notebook)
        notebook.add(performance_tab, text="Performance")
        
        # Add preference controls here...
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=10)
        ttk.Button(button_frame, text="Save", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def show_advanced_settings(self):
        """Show advanced settings dialog"""
        messagebox.showinfo("Advanced Settings", 
                          "Advanced settings allow fine-tuning conversion parameters.\n"
                          "This feature is coming soon!")
    
    def show_documentation(self):
        """Show documentation"""
        messagebox.showinfo("Documentation", 
                          "Universal Document Converter Complete\n\n"
                          "Convert documents and images to Markdown format.\n\n"
                          "Supported formats:\n"
                          "- Documents: DOCX, PDF, TXT, RTF, ODT, HTML, EPUB, XML, JSON, CSV\n"
                          "- Images (with OCR): JPG, PNG, TIFF, BMP, GIF, WebP\n\n"
                          "For more information, visit the project repository.")
    
    def show_shortcuts(self):
        """Show keyboard shortcuts"""
        shortcuts = """
        Keyboard Shortcuts:
        
        Ctrl+O - Select input folder
        Ctrl+S - Select output folder  
        Ctrl+R - Start conversion
        Escape - Stop conversion
        F1 - Show help
        
        Drag & Drop:
        - Drag folders directly onto the window
        """
        messagebox.showinfo("Keyboard Shortcuts", shortcuts)
    
    def show_about(self):
        """Show about dialog"""
        about_text = """
        Universal Document Converter Complete
        Version 3.0
        
        Designed and built by Beau Lewis
        Email: blewisxx@gmail.com
        
        A comprehensive document and image conversion tool
        with OCR support and advanced features.
        
        © 2024 All rights reserved
        """
        messagebox.showinfo("About", about_text)
    
    def show_recent_folders(self):
        """Show recent folders dialog"""
        recent = self.config_manager.get('recent_folders', [])
        if not recent:
            messagebox.showinfo("Recent Folders", "No recent folders.")
            return
        
        # Create dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Recent Folders")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Create listbox
        listbox = tk.Listbox(dialog)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        for folder in recent:
            listbox.insert(tk.END, folder)
        
        def select_folder():
            selection = listbox.curselection()
            if selection:
                folder = listbox.get(selection[0])
                self.input_path.set(folder)
                self.scan_input_folder()
                dialog.destroy()
        
        # Buttons
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=10)
        ttk.Button(button_frame, text="Select", command=select_folder).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def add_to_recent_folders(self, folder):
        """Add folder to recent folders list"""
        recent = self.config_manager.get('recent_folders', [])
        if folder in recent:
            recent.remove(folder)
        recent.insert(0, folder)
        recent = recent[:10]  # Keep only last 10
        self.config_manager.set('recent_folders', recent)
    
    def on_window_configure(self, event):
        """Handle window configuration changes"""
        if event.widget == self.root:
            # Save window geometry
            self.config_manager.set('window_geometry', self.root.geometry())
    
    def on_closing(self):
        """Handle window closing"""
        # Stop any running conversion
        if self.processing_thread and self.processing_thread.is_alive():
            self.cancel_processing.set()
            self.processing_thread.join(timeout=2)
        
        # Save configuration
        self.config_manager.save_config()
        
        # Close window
        self.root.destroy()

def main():
    """Application entry point"""
    try:
        # Try enhanced drag-and-drop support
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
    except ImportError:
        # Fall back to regular tkinter
        root = tk.Tk()
    
    # Set application icon if available
    try:
        if platform.system() == 'Windows':
            root.iconbitmap('icon.ico')
    except Exception:
        pass
    
    # Create application
    app = DocumentConverterApp(root)
    
    # Center window
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    
    # Start application
    root.mainloop()

if __name__ == "__main__":
    main()